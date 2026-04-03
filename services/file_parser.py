"""
File Parser Service
-------------------
Responsible for extracting text from PDF and DOCX files.
This module knows NOTHING about Flask, routes, or web requests.
It just takes a file path and returns text. That's it.
"""

import fitz  # This is PyMuPDF — confusing, but that's the import name
from docx import Document


def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file using PyMuPDF.
    
    Returns a dictionary with:
    - 'text': the full extracted text
    - 'pages': list of text per page (useful for page references later)
    - 'page_count': total number of pages
    - 'is_scanned': True if the PDF has little/no extractable text
    
    Why return a dictionary instead of just a string?
    Because later we'll need page numbers for references (FR-25),
    and we need to detect scanned PDFs (FR-04).
    """
    try:
        # fitz.open() opens the PDF and gives us a document object
        doc = fitz.open(file_path)
        
        pages = []
        full_text = ""
        
        for page_num in range(len(doc)):
            # doc[page_num] gets a single page
            page = doc[page_num]
            
            # get_text() extracts all text from that page
            page_text = page.get_text()
            
            pages.append({
                'page_number': page_num + 1,  # +1 because humans count from 1, not 0
                'text': page_text
            })
            
            full_text += page_text + "\n"
        
        doc.close()
        
        # --- Detect scanned PDFs (FR-04) ---
        # If a 100-page PDF has barely any text, it's probably scanned images
        # We check: is the average text per page suspiciously low?
        avg_chars_per_page = len(full_text.strip()) / max(len(pages), 1)
        is_scanned = avg_chars_per_page < 50  # Less than 50 chars per page = probably scanned
        
        return {
            'text': full_text.strip(),
            'pages': pages,
            'page_count': len(pages),
            'is_scanned': is_scanned
        }
    
    except Exception as e:
        # NFR-05: Handle corrupt/malformed files without crashing
        return {
            'text': '',
            'pages': [],
            'page_count': 0,
            'is_scanned': False,
            'error': f'Failed to parse PDF: {str(e)}'
        }


def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file using python-docx.
    
    DOCX files don't have "pages" in the same way PDFs do —
    page breaks depend on the printer/viewer settings.
    So we extract paragraphs instead, which maps better to
    document structure (headings, sections, etc.)
    """
    try:
        # Document() loads the DOCX file
        doc = Document(file_path)
        
        paragraphs = []
        full_text = ""
        
        for para in doc.paragraphs:
            # Skip empty paragraphs
            if para.text.strip():
                paragraphs.append({
                    'text': para.text.strip(),
                    # para.style.name tells us if it's a heading, normal text, etc.
                    # This will be useful later for section detection (FR-36)
                    'style': para.style.name
                })
                full_text += para.text.strip() + "\n"
        
        return {
            'text': full_text.strip(),
            'paragraphs': paragraphs,
            'paragraph_count': len(paragraphs),
            'is_scanned': False  # DOCX files are always text-based
        }
    
    except Exception as e:
        return {
            'text': '',
            'paragraphs': [],
            'paragraph_count': 0,
            'is_scanned': False,
            'error': f'Failed to parse DOCX: {str(e)}'
        }


def extract_text(file_path):
    """
    Main entry point — detects file type and calls the right extractor.
    
    This is the function other parts of the app will call.
    They don't need to know whether it's a PDF or DOCX —
    they just call extract_text() and get back the result.
    
    This pattern is called a "facade" — a simple interface
    that hides the complexity behind it.
    """
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        return {
            'text': '',
            'error': f'Unsupported file type: {file_path}'
        }
    
    